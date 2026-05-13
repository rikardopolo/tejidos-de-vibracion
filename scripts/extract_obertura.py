#!/usr/bin/env python3
"""
extract_obertura.py
Extrae Obertura_TDV_v2_09-05-2026.docx a MDX + copia figuras optimizadas.
"""
import os
import re
import sys
import textwrap
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

# ─── Rutas ───────────────────────────────────────────────────────────────────
DOCX_PATH  = Path("H:/Mi unidad/TDV-Workspace/02-Libro/capitulos-borradores/Obertura/Obertura_TDV_v2_09-05-2026.docx")
FIGS_SRC   = Path("H:/Mi unidad/TDV-Workspace/07-Assets/imagenes/Obertura")
FIGS_DST   = Path("E:/dev/tejidos-de-vibracion/public/assets/obertura")
MDX_OUT    = Path("E:/dev/tejidos-de-vibracion/src/content/book/obertura.mdx")

FIGS_DST.mkdir(parents=True, exist_ok=True)

# ─── Mapeo de figuras ─────────────────────────────────────────────────────────
FIG_MAP = {
    "O.1": ("Figura O.1.png", "figura-o-1.png"),
    "O.2": ("Figura O.2.png", "figura-o-2.png"),
    "O.3": ("Figura O.3.png", "figura-o-3.png"),
    "O.4": ("Figura O.4.png", "figura-o-4.png"),
    "O.5": ("Figura O.5.png", "figura-o-5.png"),
}

# ─── Copiar y optimizar figuras ───────────────────────────────────────────────
fig_dims = {}
for key, (src_name, dst_name) in FIG_MAP.items():
    src = FIGS_SRC / src_name
    dst = FIGS_DST / dst_name
    if src.exists():
        img = Image.open(src)
        w, h = img.size
        if w > 1600:
            ratio = 1600 / w
            w, h = 1600, int(h * ratio)
            img = img.resize((w, h), Image.LANCZOS)
        img.save(dst, format="PNG", optimize=True)
        fig_dims[key] = (w, h)
        print(f"  OK {dst_name} ({w}x{h})")
    else:
        print(f"  ERROR No encontrado: {src}", file=sys.stderr)
        fig_dims[key] = (800, 600)

# ─── Helpers de texto ─────────────────────────────────────────────────────────
def get_para_style(para):
    """Retorna el estilo del párrafo."""
    return para.style.name if para.style else "Normal"

def runs_to_md(para):
    """Convierte los runs de un párrafo a markdown inline."""
    result = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        bold   = run.bold
        italic = run.italic
        if bold and italic:
            text = f"***{text}***"
        elif bold:
            text = f"**{text}**"
        elif italic:
            text = f"*{text}*"
        result.append(text)
    return "".join(result)

def style_to_heading(style_name):
    """Convierte nombre de estilo a nivel de heading (0 = no heading)."""
    s = style_name.lower()
    if "heading 1" in s or s == "título 1":
        return 1
    if "heading 2" in s or s == "título 2":
        return 2
    if "heading 3" in s or s == "título 3":
        return 3
    if "heading 4" in s or s == "título 4":
        return 4
    return 0

# Estilos TDV especiales → componentes Astro
SPECIAL_CONTAINERS = {
    "pausa reflexiva": "PausaReflexiva",
    "pausa científica": "PausaCientifica",
    "pausa cientifica": "PausaCientifica",
    "voz del tejido": "VozTejido",
    "voz del tejedor": "VozTejido",
    "invocación": "Invocacion",
    "invocacion": "Invocacion",
}

# ─── Detección de referencias a figuras en texto ──────────────────────────────
FIG_REF_RE = re.compile(r'[Ff]igura\s+(O\.\d+)', re.IGNORECASE)

def make_figure_mdx(key, caption=""):
    """Genera el bloque MDX de figura."""
    dst_name = FIG_MAP[key][1]
    w, h = fig_dims.get(key, (800, 600))
    cap_text = f"Figura {key}" + (f" · {caption}" if caption else "")
    return (
        f'\n<figure class="figura-libro">\n'
        f'  <img src="/assets/obertura/{dst_name}" alt="Figura {key}" '
        f'loading="lazy" width="{w}" height="{h}" />\n'
        f'  <figcaption>{cap_text}</figcaption>\n'
        f'</figure>\n'
    )

# ─── Extracción de tablas TDV ────────────────────────────────────────────────
def table_to_mdx(table):
    """
    Intenta reconocer el tipo de tabla TDV y emite el componente correspondiente.
    Las tablas TDV suelen tener 1-2 columnas:
      Col 0: etiqueta de tipo ("Pausa Reflexiva", etc.)
      Col 1: contenido
    O 1 columna donde la primera celda es el tipo.
    """
    # Recoger todo el texto de la tabla para inspección
    all_text = []
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text.strip())

    full_text = " ".join(all_text).lower()

    # Detectar tipo
    comp = None
    for key, comp_name in SPECIAL_CONTAINERS.items():
        if key in full_text:
            comp = comp_name
            break

    if not comp:
        # Tabla genérica: emitir como bloque de cita
        lines = []
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_texts:
                lines.append(" | ".join(row_texts))
        return "\n> " + "\n> ".join(lines) + "\n"

    # Recopilar contenido (ignorar la celda que solo tiene la etiqueta de tipo)
    content_parts = []
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if not cell_text:
                continue
            # Saltar si es solo la etiqueta
            skip = any(k in cell_text.lower() for k in SPECIAL_CONTAINERS.keys())
            if skip and len(cell_text) < 50:
                continue
            content_parts.append(cell_text)

    content = "\n\n".join(content_parts)
    # Limpiar leading/trailing whitespace
    content = content.strip()

    if comp == "VozTejido":
        # VozTejido tiene atributo 'cita' y 'firma'
        lines = content.split("\n")
        cita = lines[0].strip() if lines else ""
        firma = lines[-1].strip() if len(lines) > 1 else ""
        return f'\n<VozTejido cita="{cita}" firma="{firma}" />\n'

    return f'\n<{comp}>\n{content}\n</{comp}>\n'

# ─── Parser principal ─────────────────────────────────────────────────────────
doc = Document(DOCX_PATH)
body = doc.element.body

sections = []  # lista de strings de markdown
figures_inserted = set()
pending_fig_captions = {}  # key → caption text

# Primera pasada: buscar leyendas de figuras (párrafos que son solo "Figura O.X · ...")
for para in doc.paragraphs:
    text = para.text.strip()
    m = re.match(r'^[Ff]igura\s+(O\.\d+)\s*[·\-–—:]\s*(.+)$', text)
    if m:
        key = m.group(1).upper().replace("O", "O")  # normalizar
        key = m.group(1)  # ej: "O.1"
        caption = m.group(2).strip()
        pending_fig_captions[key] = caption
        print(f"  Caption encontrada: Figura {key} -> {caption[:60]}")

# Segunda pasada: extraer contenido
current_table_type = None
i = 0
children = list(body.iterchildren())
para_idx = 0

# Usaremos doc.paragraphs y doc.tables con tracking de posición
# Mejor: iterar el XML directamente
lines = []

def flush_fig(key):
    """Inserta figura si no fue insertada aun."""
    if key not in figures_inserted:
        figures_inserted.add(key)
        cap = pending_fig_captions.get(key, "")
        lines.append(make_figure_mdx(key, cap))

paras_iter = iter(doc.paragraphs)
tables_seen = 0

for child in children:
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

    if tag == "p":
        # Párrafo
        # Buscar el objeto Paragraph correspondiente
        try:
            para = next(paras_iter)
        except StopIteration:
            break

        style = get_para_style(para)
        text = para.text.strip()
        md_text = runs_to_md(para)

        # ¿Es una referencia de figura (leyenda)?
        m_cap = re.match(r'^[Ff]igura\s+(O\.\d+)', text)
        if m_cap:
            # Es el caption inline → ya capturado arriba, emitir figura aquí
            key = m_cap.group(1)
            flush_fig(key)
            continue

        # ¿Contiene referencia a figura dentro del texto?
        refs = FIG_REF_RE.findall(text)
        for ref in refs:
            flush_fig(ref)

        # Heading
        hlevel = style_to_heading(style)
        if hlevel:
            prefix = "#" * hlevel
            lines.append(f"\n{prefix} {text}\n")
            continue

        # Párrafo vacío
        if not text:
            lines.append("")
            continue

        # Párrafo de cuerpo normal
        # ¿Tiene clase especial?
        style_lower = style.lower()
        if "list" in style_lower or "bullet" in style_lower:
            lines.append(f"- {md_text}")
        elif "quote" in style_lower or "epígrafe" in style_lower:
            lines.append(f"\n> {md_text}\n")
        else:
            lines.append(f"\n{md_text}\n")

    elif tag == "tbl":
        # Tabla
        try:
            table = doc.tables[tables_seen]
            tables_seen += 1
        except IndexError:
            continue
        mdx_block = table_to_mdx(table)
        lines.append(mdx_block)

# Insertar figuras que aun no se han emitido al final (si las hay)
for key in FIG_MAP:
    if key not in figures_inserted:
        print(f"  WARN Figura {key} no encontrada en flujo, se agrega al final antes del corte")
        lines.append(make_figure_mdx(key, pending_fig_captions.get(key, "")))

# ─── Contar palabras hasta puntos de corte ────────────────────────────────────
full_content = "\n".join(lines)

# Limpiar contenido: quitar tags XML/Astro para contar palabras
def count_words(text):
    clean = re.sub(r'<[^>]+>', ' ', text)
    clean = re.sub(r'\{[^}]+\}', ' ', clean)
    words = clean.split()
    return len(words)

# Buscar secciones candidatas para el corte
section_starts = []
for i, line in enumerate(lines):
    if re.match(r'^#{1,3} ', line.strip()):
        heading = line.strip()
        words_so_far = count_words("\n".join(lines[:i]))
        section_starts.append((i, heading, words_so_far))
        print(f"  Sección en línea {i}: {heading[:60]} → {words_so_far} palabras antes")

# Elegir punto de corte más cercano a 1700 palabras
best_idx = None
best_diff = float('inf')
for idx, heading, words in section_starts:
    diff = abs(words - 1700)
    if diff < best_diff:
        best_diff = diff
        best_idx = idx
        best_heading = heading
        best_words = words

if best_idx is not None:
    print(f"\n  → Corte elegido antes de: {best_heading}")
    print(f"    Palabras antes del corte: {best_words}")
    # Insertar CorteTejedor ANTES del heading
    lines.insert(best_idx, "\n<CorteTejedor />\n")
else:
    # Fallback: insertar a mitad del contenido
    mid = len(lines) // 2
    lines.insert(mid, "\n<CorteTejedor />\n")
    print("  ! Usando corte de fallback a mitad del documento")

# ─── Contar palabras del fragmento visible (antes del corte) ─────────────────
cut_pos = None
for i, line in enumerate(lines):
    if "<CorteTejedor" in line:
        cut_pos = i
        break

if cut_pos is not None:
    fragment = "\n".join(lines[:cut_pos])
    visible_words = count_words(fragment)
    print(f"\n  Palabras visibles para Nivel 0: {visible_words}")

# ─── Ensamble del MDX ────────────────────────────────────────────────────────
total_words = count_words("\n".join(lines))
reading_time = max(1, round(total_words / 200))

FRONTMATTER = f"""---
title: "Obertura"
subtitle: "El tejedor consciente despierta"
order: 0
volume: 1
act: "obertura"
chapter: null
kind: "obertura"
status: "published"
publishedAt: 2026-04-19
authors:
  - "Ricardo Polo"
  - "Orion"
tags:
  - "obertura"
  - "tejedor-consciente"
  - "frontispicio"
readingTime: {reading_time}
---

import Flourish from '@/components/book/Flourish.astro';
import Pullquote from '@/components/book/Pullquote.astro';
import Invocacion from '@/components/book/Invocacion.astro';
import PausaReflexiva from '@/components/book/PausaReflexiva.astro';
import PausaCientifica from '@/components/book/PausaCientifica.astro';
import VozTejido from '@/components/book/VozTejido.astro';
import LitLink from '@/components/book/LitLink.astro';
import CorteTejedor from '@/components/book/CorteTejedor.astro';

"""

body_content = "\n".join(lines)
# Limpiar líneas vacías excesivas (> 2 consecutivas)
body_content = re.sub(r'\n{4,}', '\n\n\n', body_content)

mdx_final = FRONTMATTER + body_content.strip() + "\n"

MDX_OUT.write_text(mdx_final, encoding="utf-8")
print(f"\nMDX escrito: {MDX_OUT}")
print(f"  Total palabras: {total_words}")
print(f"  Reading time: {reading_time} min")
print(f"  Palabras visibles nivel 0: {visible_words if cut_pos else 'N/A'}")

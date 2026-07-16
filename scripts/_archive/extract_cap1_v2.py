#!/usr/bin/env python3
"""
extract_cap1_v2.py
Extrae Cap1 DOCX (v2_1) a MDX con deteccion de secciones por patron §N.N
y corte al final de la seccion 1.2 (~2500 palabras).
"""
import re
import sys
from pathlib import Path

from docx import Document
from PIL import Image

# ─── Rutas ───────────────────────────────────────────────────────────────────
DOCX_PATH = Path("H:/Mi unidad/TDV-Workspace/02-Libro/capitulos-borradores/Capitulo 1/Cap1_TDV_VolumenI_Modo_B3_consolidado_v2_1_09-05-2026.docx")
FIGS_SRC  = Path("H:/Mi unidad/TDV-Workspace/07-Assets/imagenes/Capitulo 1")
FIGS_DST  = Path("E:/dev/tejidos-de-vibracion/public/assets/cap-1")
MDX_OUT   = Path("E:/dev/tejidos-de-vibracion/src/content/book/cap-1-universo-sinfonia.mdx")

FIGS_DST.mkdir(parents=True, exist_ok=True)

# ─── Figuras ─────────────────────────────────────────────────────────────────
FIG_MAP = {
    "1.1": ("Figura 1.1.png", "figura-1-1.png"),
    "1.2": ("Figura 1.2.png", "figura-1-2.png"),
    "1.3": ("Figura 1.3.png", "figura-1-3.png"),
    "1.4": ("Figura 1.4.png", "figura-1-4.png"),
    "1.5": ("Figura 1.5.png", "figura-1-5.png"),
    "1.6": ("Figura 1.6.png", "figura-1-6.png"),
    "1.7": ("Figura 1.7.png", "figura-1-7.png"),
    "1.8": ("Figura 1.8.png", "figura-1-8.png"),
    "1.9": ("Figura 1.9.png", "figura-1-9.png"),
}

fig_dims = {}
for key, (src_name, dst_name) in FIG_MAP.items():
    src = FIGS_SRC / src_name
    dst = FIGS_DST / dst_name
    if src.exists():
        img = Image.open(src)
        w, h = img.size
        if w > 1600:
            ratio = 1600 / w
            w, h = 1600, int(h * ratio)
            img = img.resize((w, h), Image.LANCZOS)
        img.save(dst, format="PNG", optimize=True)
        fig_dims[key] = (w, h)
        print(f"  OK {dst_name} ({w}x{h})")
    else:
        print(f"  ERROR no encontrado: {src}", file=sys.stderr)
        fig_dims[key] = (800, 600)

# ─── Helpers ──────────────────────────────────────────────────────────────────
FIG_CAP_RE  = re.compile(r'^[Ff]igura\s+(1\.\d+)\s*[·\-–—:]\s*(.+)$')
FIG_REF_RE  = re.compile(r'[Ff]igura\s+(1\.\d+)', re.IGNORECASE)
SEC_RE      = re.compile(r'^§(\d+\.\d+)\s*[·\-–—:]\s*(.+)$')
FORM_RE     = re.compile(r'^[Ff]órmula\s+\d+')
ORNAMENT_RE = re.compile(r'^[◆◇\s]+$')

SPECIAL_CONTAINERS = {
    "pausa reflexiva": "PausaReflexiva",
    "pausa científica": "PausaCientifica",
    "pausa cientifica": "PausaCientifica",
    "voz del tejido": "VozTejido",
    "voz del tejedor": "VozTejido",
    "invocación": "Invocacion",
    "invocacion": "Invocacion",
}

def runs_to_md(para):
    result = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        b, i = run.bold, run.italic
        if b and i: text = f"***{text}***"
        elif b:     text = f"**{text}**"
        elif i:     text = f"*{text}*"
        result.append(text)
    return "".join(result)

def make_figure_mdx(key, caption=""):
    dst_name = FIG_MAP[key][1]
    w, h = fig_dims.get(key, (800, 600))
    cap_text = f"Figura {key}" + (f" · {caption}" if caption else "")
    return (
        f'\n<figure class="figura-libro">\n'
        f'  <img src="/assets/cap-1/{dst_name}" alt="Figura {key}" '
        f'loading="lazy" width="{w}" height="{h}" />\n'
        f'  <figcaption>{cap_text}</figcaption>\n'
        f'</figure>\n'
    )

def table_to_mdx(table):
    all_text = " ".join(
        cell.text.strip()
        for row in table.rows
        for cell in row.cells
    ).lower()

    comp = None
    for key, comp_name in SPECIAL_CONTAINERS.items():
        if key in all_text:
            comp = comp_name
            break

    if not comp:
        lines = []
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_texts:
                lines.append(" | ".join(row_texts))
        return ("\n> " + "\n> ".join(lines) + "\n") if lines else ""

    content_parts = []
    for row in table.rows:
        for cell in row.cells:
            ct = cell.text.strip()
            if not ct:
                continue
            if any(k in ct.lower() for k in SPECIAL_CONTAINERS) and len(ct) < 60:
                continue
            content_parts.append(ct)
    content = "\n\n".join(content_parts).strip()

    if comp == "VozTejido":
        ls = content.split("\n")
        cita = ls[0].strip() if ls else ""
        firma = ls[-1].strip() if len(ls) > 1 else ""
        return f'\n<VozTejido cita="{cita}" firma="{firma}" />\n'

    return f'\n<{comp}>\n{content}\n</{comp}>\n'

# ─── Extracción ───────────────────────────────────────────────────────────────
doc = Document(DOCX_PATH)

# Primera pasada: recoger captions de figuras
pending_caps = {}
for para in doc.paragraphs:
    text = para.text.strip()
    m = FIG_CAP_RE.match(text)
    if m:
        pending_caps[m.group(1)] = m.group(2).strip()
        print(f"  Caption Figura {m.group(1)} -> {m.group(2)[:60]}")

# Segunda pasada: extraer contenido
doc2 = Document(DOCX_PATH)
figures_inserted = set()
lines = []

# Detectar inicio del contenido
SKIP_KEYWORDS = [
    "capítulo  1", "capítulo 1",
    "la historia secreta del sonido",
    "modo b3", "consolidado", "tejidos de realidad",
    "tejidosderealidad", "safecreative", "isbn",
    "© 2026", "primera edición", "prohibida",
    "cámara colombiana", "agencia isbn",
    "como tecnología de conciencia",
]

content_started = False
paras = doc2.paragraphs

def flush_fig(key):
    if key not in figures_inserted and key in FIG_MAP:
        figures_inserted.add(key)
        lines.append(make_figure_mdx(key, pending_caps.get(key, "")))

tables_seen = 0
from docx.oxml.ns import qn
body = doc2.element.body
paras_iter = iter(doc2.paragraphs)

for child in body.iterchildren():
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

    if tag == "p":
        try:
            para = next(paras_iter)
        except StopIteration:
            break

        text = para.text.strip()
        md_text = runs_to_md(para)

        # Detectar inicio: el "Umbral Poético" del capítulo
        if not content_started:
            if text == "Umbral Poético" or text == "Umbral poético":
                content_started = True
                lines.append("\n## Umbral poético\n")
                continue
            # También empezar si encontramos la primera sección §
            if SEC_RE.match(text):
                content_started = True
            else:
                continue

        if not text:
            lines.append("")
            continue

        # Ornamentos
        if ORNAMENT_RE.match(text):
            if "◇" in text and "◆" in text:
                lines.append("\n<Flourish kind=\"section\" />\n")
            else:
                lines.append("\n<Flourish kind=\"diamond\" />\n")
            continue

        # Caption de figura → emit figura aquí
        m_cap = FIG_CAP_RE.match(text)
        if m_cap:
            flush_fig(m_cap.group(1))
            continue

        # Referencias inline a figuras
        for ref in FIG_REF_RE.findall(text):
            flush_fig(ref)

        # Sección §N.N
        m_sec = SEC_RE.match(text)
        if m_sec:
            sec_num = m_sec.group(1)
            sec_title = m_sec.group(2).strip()
            lines.append(f"\n## {sec_num} · {sec_title}\n")
            continue

        # Fórmula matemática
        if FORM_RE.match(text):
            lines.append(f"\n*{text}*\n")
            continue

        # Líneas de fórmula (simples, como "E = h · ν")
        if re.match(r'^[A-Za-zÀ-ú]\s*=\s*', text) and len(text) < 30:
            lines.append(f"\n```\n{text}\n```\n")
            continue

        # Párrafo normal
        lines.append(f"\n{md_text}\n")

    elif tag == "tbl":
        try:
            table = doc2.tables[tables_seen]
            tables_seen += 1
        except IndexError:
            continue
        if content_started:
            block = table_to_mdx(table)
            if block:
                lines.append(block)

# ─── Determinar posicion del CorteTejedor ────────────────────────────────────
# Corte justo antes de §1.3 (aprox 2500 palabras del cuerpo)
cut_idx = None
for i, line in enumerate(lines):
    if "## 1.3" in line:
        cut_idx = i
        break

if cut_idx is not None:
    lines.insert(cut_idx, "\n<CorteTejedor />\n")
    print(f"\nCorteTejedor insertado antes de seccion 1.3 (linea {cut_idx})")
else:
    # Fallback: ~2500 palabras
    words_so_far = 0
    for i, line in enumerate(lines):
        clean = re.sub(r'<[^>]+>', ' ', line)
        words_so_far += len(clean.split())
        if words_so_far >= 2400:
            lines.insert(i+1, "\n<CorteTejedor />\n")
            print(f"  CorteTejedor fallback en linea {i+1}, palabras: {words_so_far}")
            break

# ─── Insertar figuras no encontradas DESPUES del corte ───────────────────────
# Figuras 1.1, 1.2, 1.3 no tienen referencia en texto → van al final (Nivel 1+)
for fig_key in ["1.1", "1.2", "1.3", "1.6"]:
    if fig_key not in figures_inserted and fig_key in FIG_MAP:
        print(f"  INFO Figura {fig_key} sin referencia en texto, agrega al final")
        figures_inserted.add(fig_key)
        lines.append(make_figure_mdx(fig_key, pending_caps.get(fig_key, "")))

# ─── Palabras visibles ────────────────────────────────────────────────────────
cut_pos = None
for i, line in enumerate(lines):
    if "<CorteTejedor" in line:
        cut_pos = i
        break

def count_words(text):
    clean = re.sub(r'<[^>]+>', ' ', text)
    clean = re.sub(r'\{[^}]+\}', ' ', clean)
    return len(clean.split())

visible_words = count_words("\n".join(lines[:cut_pos])) if cut_pos else 0
print(f"Palabras visibles nivel 0: {visible_words}")

# ─── Ensamble ─────────────────────────────────────────────────────────────────
full_body = "\n".join(lines)
full_body = re.sub(r'\n{4,}', '\n\n\n', full_body)

total_words = count_words(full_body)
reading_time = max(1, round(total_words / 200))

FRONTMATTER = f"""---
title: "La historia secreta del sonido"
subtitle: "Como tecnología de conciencia"
order: 1
volume: 1
act: "I"
chapter: 1
kind: "capitulo"
status: "published"
publishedAt: 2026-05-01
authors:
  - "Ricardo Polo"
tags:
  - "sonido"
  - "vibración"
  - "cymática"
  - "consciencia"
  - "AUM"
readingTime: {reading_time}
---

import Flourish from '@/components/book/Flourish.astro';
import Pullquote from '@/components/book/Pullquote.astro';
import Invocacion from '@/components/book/Invocacion.astro';
import PausaReflexiva from '@/components/book/PausaReflexiva.astro';
import PausaCientifica from '@/components/book/PausaCientifica.astro';
import VozTejido from '@/components/book/VozTejido.astro';
import LitLink from '@/components/book/LitLink.astro';
import CorteTejedor from '@/components/book/CorteTejedor.astro';

"""

mdx_final = FRONTMATTER + full_body.strip() + "\n"
MDX_OUT.write_text(mdx_final, encoding="utf-8")

print(f"\nMDX escrito: {MDX_OUT}")
print(f"  Total palabras: {total_words}")
print(f"  Reading time: {reading_time} min")
print(f"  Palabras visibles nivel 0: {visible_words}")

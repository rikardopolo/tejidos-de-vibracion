#!/usr/bin/env python3
"""
extract_cap1.py
Extrae Cap1_TDV_VolumenI_Modo_B3_consolidado_v2_1_09-05-2026.docx a MDX
+ copia 9 figuras optimizadas.
"""
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

# ─── Rutas ───────────────────────────────────────────────────────────────────
DOCX_PATH = Path("H:/Mi unidad/TDV-Workspace/02-Libro/capitulos-borradores/Capitulo 1/Cap1_TDV_VolumenI_Modo_B3_consolidado_v2_1_09-05-2026.docx")
FIGS_SRC  = Path("H:/Mi unidad/TDV-Workspace/07-Assets/imagenes/Capitulo 1")
FIGS_DST  = Path("E:/dev/tejidos-de-vibracion/public/assets/cap-1")
MDX_OUT   = Path("E:/dev/tejidos-de-vibracion/src/content/book/cap-1-universo-sinfonia.mdx")

FIGS_DST.mkdir(parents=True, exist_ok=True)

# ─── Mapeo de figuras ─────────────────────────────────────────────────────────
FIG_MAP = {
    "1.1": ("Figura 1.1.png", "figura-1-1.png"),
    "1.2": ("Figura 1.2.png", "figura-1-2.png"),
    "1.3": ("Figura 1.3.png", "figura-1-3.png"),
    "1.4": ("Figura 1.4.png", "figura-1-4.png"),
    "1.5": ("Figura 1.5.png", "figura-1-5.png"),
    "1.6": ("Figura 1.6.png", "figura-1-6.png"),
    "1.7": ("Figura 1.7.png", "figura-1-7.png"),
    "1.8": ("Figura 1.8.png", "figura-1-8.png"),
    "1.9": ("Figura 1.9.png", "figura-1-9.png"),
}

# ─── Copiar y optimizar figuras ───────────────────────────────────────────────
fig_dims = {}
for key, (src_name, dst_name) in FIG_MAP.items():
    src = FIGS_SRC / src_name
    dst = FIGS_DST / dst_name
    if src.exists():
        img = Image.open(src)
        w, h = img.size
        if w > 1600:
            ratio = 1600 / w
            w, h = 1600, int(h * ratio)
            img = img.resize((w, h), Image.LANCZOS)
        img.save(dst, format="PNG", optimize=True)
        fig_dims[key] = (w, h)
        print(f"  OK {dst_name} ({w}x{h})")
    else:
        print(f"  ERROR No encontrado: {src}", file=sys.stderr)
        fig_dims[key] = (800, 600)

# ─── Helpers ──────────────────────────────────────────────────────────────────
def get_para_style(para):
    return para.style.name if para.style else "Normal"

def runs_to_md(para):
    result = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        bold   = run.bold
        italic = run.italic
        if bold and italic:
            text = f"***{text}***"
        elif bold:
            text = f"**{text}**"
        elif italic:
            text = f"*{text}*"
        result.append(text)
    return "".join(result)

def style_to_heading(style_name):
    s = style_name.lower()
    if "heading 1" in s or "título 1" in s:
        return 1
    if "heading 2" in s or "título 2" in s:
        return 2
    if "heading 3" in s or "título 3" in s:
        return 3
    if "heading 4" in s or "título 4" in s:
        return 4
    return 0

SPECIAL_CONTAINERS = {
    "pausa reflexiva": "PausaReflexiva",
    "pausa científica": "PausaCientifica",
    "pausa cientifica": "PausaCientifica",
    "voz del tejido": "VozTejido",
    "voz del tejedor": "VozTejido",
    "invocación": "Invocacion",
    "invocacion": "Invocacion",
}

FIG_REF_RE = re.compile(r'[Ff]igura\s+(1\.\d+)', re.IGNORECASE)

def make_figure_mdx(key, caption=""):
    dst_name = FIG_MAP[key][1]
    w, h = fig_dims.get(key, (800, 600))
    cap_text = f"Figura {key}" + (f" · {caption}" if caption else "")
    return (
        f'\n<figure class="figura-libro">\n'
        f'  <img src="/assets/cap-1/{dst_name}" alt="Figura {key}" '
        f'loading="lazy" width="{w}" height="{h}" />\n'
        f'  <figcaption>{cap_text}</figcaption>\n'
        f'</figure>\n'
    )

def table_to_mdx(table):
    all_text = []
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text.strip())
    full_text = " ".join(all_text).lower()

    comp = None
    for key, comp_name in SPECIAL_CONTAINERS.items():
        if key in full_text:
            comp = comp_name
            break

    if not comp:
        lines = []
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_texts:
                lines.append(" | ".join(row_texts))
        if lines:
            return "\n> " + "\n> ".join(lines) + "\n"
        return ""

    content_parts = []
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if not cell_text:
                continue
            skip = any(k in cell_text.lower() for k in SPECIAL_CONTAINERS.keys())
            if skip and len(cell_text) < 60:
                continue
            content_parts.append(cell_text)

    content = "\n\n".join(content_parts).strip()

    if comp == "VozTejido":
        lines = content.split("\n")
        cita = lines[0].strip() if lines else ""
        firma = lines[-1].strip() if len(lines) > 1 else ""
        return f'\n<VozTejido cita="{cita}" firma="{firma}" />\n'

    return f'\n<{comp}>\n{content}\n</{comp}>\n'

# ─── Extracción ───────────────────────────────────────────────────────────────
doc = Document(DOCX_PATH)
body = doc.element.body

# Primera pasada: capturar leyendas de figuras
pending_fig_captions = {}
for para in doc.paragraphs:
    text = para.text.strip()
    m = re.match(r'^[Ff]igura\s+(1\.\d+)\s*[·\-–—:]\s*(.+)$', text)
    if m:
        key = m.group(1)
        caption = m.group(2).strip()
        pending_fig_captions[key] = caption
        print(f"  Caption: Figura {key} -> {caption[:60]}")

# Segunda pasada: extraer contenido
lines = []
figures_inserted = set()
paras_iter = iter(doc.paragraphs)
tables_seen = 0
children = list(body.iterchildren())
content_started = False
skip_until_content = True

def flush_fig(key):
    if key not in figures_inserted and key in FIG_MAP:
        figures_inserted.add(key)
        cap = pending_fig_captions.get(key, "")
        lines.append(make_figure_mdx(key, cap))

# Palabras clave que indican inicio del contenido real
CONTENT_START_MARKERS = [
    "umbral poético", "umbral poetico",
    "hay un instante", "hay una frecuencia",
    "capítulo 1", "capitulo 1",
    "1.1", "1. ",
]

for child in children:
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

    if tag == "p":
        try:
            para = next(paras_iter)
        except StopIteration:
            break

        style = get_para_style(para)
        text = para.text.strip()
        md_text = runs_to_md(para)

        # Detectar inicio del contenido (saltar portada)
        if skip_until_content:
            text_lower = text.lower()
            is_heading = style_to_heading(style) > 0
            # Empezar cuando encontramos heading o texto de contenido
            if is_heading or any(m in text_lower for m in CONTENT_START_MARKERS):
                skip_until_content = False
            elif text and not any(skip_kw in text_lower for skip_kw in [
                "isbn", "copyright", "safecreative", "tejidos de realidad",
                "primera edición", "prohibida", "registrada",
                "volumen i", "2026", "©", "cámara", "isbn",
                "ricardo polo", "modo b3", "◇", "tejidosderealidad"
            ]) and len(text) > 80:
                # Párrafo largo que no es portada
                skip_until_content = False

            if skip_until_content:
                continue

        # ¿Es una referencia de figura (leyenda)?
        m_cap = re.match(r'^[Ff]igura\s+(1\.\d+)', text)
        if m_cap:
            key = m_cap.group(1)
            flush_fig(key)
            continue

        # ¿Contiene referencia a figura dentro del texto?
        refs = FIG_REF_RE.findall(text)
        for ref in refs:
            flush_fig(ref)

        hlevel = style_to_heading(style)
        if hlevel:
            prefix = "#" * hlevel
            lines.append(f"\n{prefix} {text}\n")
            continue

        if not text:
            lines.append("")
            continue

        style_lower = style.lower()
        if "list" in style_lower or "bullet" in style_lower:
            lines.append(f"- {md_text}")
        elif "quote" in style_lower or "epígrafe" in style_lower:
            lines.append(f"\n> {md_text}\n")
        else:
            lines.append(f"\n{md_text}\n")

    elif tag == "tbl":
        try:
            table = doc.tables[tables_seen]
            tables_seen += 1
        except IndexError:
            continue
        if not skip_until_content:
            mdx_block = table_to_mdx(table)
            if mdx_block:
                lines.append(mdx_block)

# Insertar figuras no encontradas
for key in FIG_MAP:
    if key not in figures_inserted:
        print(f"  WARN Figura {key} no encontrada en flujo del texto")

# ─── Encontrar secciones ─────────────────────────────────────────────────────
def count_words(text):
    clean = re.sub(r'<[^>]+>', ' ', text)
    clean = re.sub(r'\{[^}]+\}', ' ', clean)
    return len(clean.split())

section_starts = []
for i, line in enumerate(lines):
    if re.match(r'^#{1,3} ', line.strip()):
        heading = line.strip()
        words_so_far = count_words("\n".join(lines[:i]))
        section_starts.append((i, heading, words_so_far))
        print(f"  Seccion en linea {i}: {heading[:60]} -> {words_so_far} palabras")

# Elegir corte más cercano a 2500 palabras
# Preferir el corte al final de la primera sección principal numerada (1.1)
best_idx = None
best_diff = float('inf')
TARGET = 2500

for idx, heading, words in section_starts:
    diff = abs(words - TARGET)
    if diff < best_diff:
        best_diff = diff
        best_idx = idx
        best_heading = heading
        best_words = words

if best_idx is not None:
    print(f"\n  Corte elegido antes de: {best_heading}")
    print(f"    Palabras antes del corte: {best_words}")
    lines.insert(best_idx, "\n<CorteTejedor />\n")
else:
    mid = len(lines) // 2
    lines.insert(mid, "\n<CorteTejedor />\n")
    print("  WARN: Usando corte fallback")

# ─── Contar palabras visibles ─────────────────────────────────────────────────
cut_pos = None
for i, line in enumerate(lines):
    if "<CorteTejedor" in line:
        cut_pos = i
        break

if cut_pos:
    fragment = "\n".join(lines[:cut_pos])
    visible_words = count_words(fragment)
    print(f"\n  Palabras visibles Nivel 0: {visible_words}")

# ─── Ensamble MDX ────────────────────────────────────────────────────────────
full_content = "\n".join(lines)
full_content = re.sub(r'\n{4,}', '\n\n\n', full_content)

total_words = count_words(full_content)
reading_time = max(1, round(total_words / 200))

FRONTMATTER = f"""---
title: "El universo como sinfonía"
subtitle: "La frecuencia como lenguaje fundamental del cosmos"
order: 1
volume: 1
act: "I"
chapter: 1
kind: "capitulo"
status: "published"
publishedAt: 2026-05-01
authors:
  - "Ricardo Polo"
tags:
  - "frecuencia"
  - "vibración"
  - "cymática"
  - "física"
  - "sinfonía-cósmica"
readingTime: {reading_time}
---

import Flourish from '@/components/book/Flourish.astro';
import Pullquote from '@/components/book/Pullquote.astro';
import Invocacion from '@/components/book/Invocacion.astro';
import PausaReflexiva from '@/components/book/PausaReflexiva.astro';
import PausaCientifica from '@/components/book/PausaCientifica.astro';
import VozTejido from '@/components/book/VozTejido.astro';
import LitLink from '@/components/book/LitLink.astro';
import CorteTejedor from '@/components/book/CorteTejedor.astro';

"""

mdx_final = FRONTMATTER + full_content.strip() + "\n"
MDX_OUT.write_text(mdx_final, encoding="utf-8")

print(f"\nMDX escrito: {MDX_OUT}")
print(f"  Total palabras: {total_words}")
print(f"  Reading time: {reading_time} min")
print(f"  Palabras visibles nivel 0: {visible_words if cut_pos else 'N/A'}")
